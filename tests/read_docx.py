from docx import Document
import sys
sys.stdout.reconfigure(encoding="utf-8")

doc = Document(r"D:\daily-ai-news\docs\真实rss信息输入.docx")
for para in doc.paragraphs:
    print(para.text)
for table in doc.tables:
    print("\n=== TABLE ===")
    for row in table.rows:
        print(" | ".join([cell.text for cell in row.cells]))

